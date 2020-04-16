#!/usr/bin/env python
import sys
import re
import os

import json

from subprocess import call
from docxtpl import DocxTemplate, RichText, InlineImage
from docx.shared import Mm, Inches, Pt, Cm, Emu
from app.utils.tools import FilePath
#import jinja2
#from jinja2.utils import Markup

class SubDoc(object):
	def __init__(self, template, images=0):
		self._template = template
		self._subtpl = template.new_subdoc()
		self._img_path = images

	def add_cover(self):
		section = self._template.sections[0]
		section.left_margin = Emu(0)
		section.top_margin = Emu(0)
		section.bottom_margin = Emu(0)
		section.right_margin = Emu(0)
		paragraph = self._template.paragraphs[0]
		run = paragraph.add_run()
		if self._img_path is not 0:
			run.add_picture(self._img_path, height=Cm(29.7), width=Cm(21.0))

	def add_endcover(self):
		section = self._template.add_section()
		section.left_margin = Emu(0)
		section.top_margin = Emu(0)
		section.bottom_margin = Emu(0)
		section.right_margin = Emu(0)
		if self._img_path is not 0:
			self._template.add_picture(self._img_path, height=Cm(29.7), width=Cm(21.0))

class Json2Doc(object):
	"""docstring for ClassName"""
	def __init__(self, injson, outdoc, tplfile, cover=True):
		self.__FilePath = FilePath()
		self._injson = injson
		self._tplfile = tplfile
		self._cover = cover
		self.report_path = outdoc
		self.pdf_path = ''
		self.main()
		print('Docx Report has been saved!!!! Report path: {}'.format(self.report_path))
		print('PDF Report has been saved! Report path: {}'.format(self.pdf_path))

	def main(self):
		tpl = DocxTemplate(self._tplfile)
		with open(self._injson, 'r') as fi:
			info_dict = json.load(fi)

		"""get covers """
		

		if self._cover:
			"""
			###线下获取模板
			try:
				channel_name = info_dict['SampleInfo']['channel_name']
			except:
				channel_name = ''
			covers = CommonMethods.get_cover_pic(channel_name, source='Covers')
			endcovers = CommonMethods.get_cover_pic(channel_name, source='EndCovers')
			"""
			### 线上传递过来 封面信息
			try:
				covers = info_dict['SampleInfo']['template_info']['front_cover']
			except:
				covers = []
			try:
				endcovers = info_dict['SampleInfo']['template_info']['end_cover']
			except:
				covers = []
			for each in covers:
				SubDoc(tpl, each).add_cover()
			for each in endcovers:
				SubDoc(tpl, each).add_endcover()
		
		
		### 转化图片
		for item in ['indiCharacterItems','comDiseaseItems','DrugsItems']:
			try:
				for i in range(0,len(info_dict[item][0]['sub'])):
					info_dict[item][0]['sub'][i]['resultimg'] = InlineImage(tpl, info_dict[item][0]['sub'][i]['resultimg'].replace('/static','RptSer'))
			except:
				pass

		tpl.render(info_dict)
		temp_path = self.__FilePath.get_common_report_path()
		tpl.save(temp_path)
		java_path = self.__FilePath.get_word2pdf_java_path()
		self.pdf_path = self.__FilePath.get_report_pdf_path(self.report_path)
		call('java -jar {} {} {} {}'.format(java_path, temp_path, self.report_path, '--nopdf'), shell=True)
		call('java -jar {} {} {}'.format(java_path, self.report_path, self.pdf_path ), shell=True)
		os.remove(temp_path)
		
